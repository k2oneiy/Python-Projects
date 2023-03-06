

from docx import Document
from docx.shared import Inches

document = Document()

name = "IYAPPAN M"
title = "Software Developer"

document.add_picture('iya.png',width=Inches(1.0))

document.add_heading(
    name + ' | ' + title,1
)

document.add_heading('Objective',2)

document.add_paragraph('Passion become an glorious Software Developer as well as Cyber security Researcher.')


document.add_heading('Experience',2)

p = document.add_paragraph()
p.add_run('Chadura Tech'+ ' '+'\n').bold = True
p.add_run('June 2021 - December 2021'+'\n').italic = True
p.add_run('Worked With Python as well as Django')


document.add_heading('Skills',2)

p = document.add_paragraph()

p.add_run('Python'+'\n')

p.add_run('Django'+'\n')

p.add_run('FastApi'+'\n')


document.save("cv.docx")




